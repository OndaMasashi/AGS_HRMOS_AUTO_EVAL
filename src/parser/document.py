"""文書解析モジュール - PDF/Word/Excelからテキスト抽出"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """ファイルからテキストを抽出する。対応形式: PDF, DOCX, XLSX"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _extract_from_pdf(file_path)
        elif suffix == ".docx":
            return _extract_from_docx(file_path)
        elif suffix in (".xlsx", ".xls"):
            return _extract_from_xlsx(file_path)
        elif suffix == ".doc":
            logger.warning(f"古い .doc 形式は非対応です: {path.name}")
            return ""
        else:
            logger.warning(f"非対応のファイル形式: {path.name} ({suffix})")
            return ""
    except Exception as e:
        logger.error(f"テキスト抽出に失敗: {path.name} - {e}")
        return ""


def _extract_from_pdf(file_path: str) -> str:
    """PDFからテキストを抽出"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n".join(text_parts)
    logger.info(f"  PDF解析完了: {len(text)} 文字抽出")
    return text


def _extract_from_docx(file_path: str) -> str:
    """Word (.docx) からテキストを抽出"""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    # 本文パラグラフ
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # テーブル内のテキストも抽出
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    text = "\n".join(text_parts)
    logger.info(f"  DOCX解析完了: {len(text)} 文字抽出")
    return text


def _extract_from_xlsx(file_path: str) -> str:
    """Excel (.xlsx) からテキストを抽出"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            for cell_value in row:
                if cell_value is not None:
                    text_parts.append(str(cell_value))

    text = "\n".join(text_parts)
    logger.info(f"  XLSX解析完了: {len(text)} 文字抽出")
    return text
